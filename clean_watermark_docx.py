import json
import os
import tempfile
import zipfile

from docx import Document
from docx.shared import Inches
from lxml import etree


SIZE_THRESHOLD = Inches(0.2)
HEADER_FOOTER_KEYWORDS = ["学科网", "zxxk.com", "北京）股份有限公司"]
BODY_TEXT_KEYWORDS = [
    "zxxk.com",
    "学科网（北京）股份有限公司",
    "学科网(北京)股份有限公司",
    "北京）股份有限公司",
]
METADATA_KEYWORDS = ["学科网", "zxxk.com", "rbm.xkw.com", "xkw"]
BODY_DRAWING_KEYWORDS = ["学科网", "zxxk.com", "rbm.xkw.com", "xkw"]
CUSTOM_PROPERTY_NS = {
    "cp": "http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"
}
CONFIG_FILENAME = "config.json"
DEFAULT_CONFIG = {
    "metadata_keywords": METADATA_KEYWORDS,
    "docx_core_properties": {
        "override_enabled": False,
        "values": {
            "author": "User",
            "comments": "",
            "title": "清理后的文档",
            "subject": "",
            "keywords": "",
        },
    },
}


def build_output_path(input_path):
    root, _ = os.path.splitext(input_path)
    return f"{root}(cleaned).docx"


def has_keyword_text(text, keywords):
    return any(keyword.lower() in text.lower() for keyword in keywords)


def first_matched_keyword(text, keywords):
    lower_text = text.lower()
    for keyword in keywords:
        if keyword.lower() in lower_text:
            return keyword
    return None


def load_config():
    config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), CONFIG_FILENAME)
    config = {
        "metadata_keywords": list(DEFAULT_CONFIG["metadata_keywords"]),
        "docx_core_properties": {
            "override_enabled": DEFAULT_CONFIG["docx_core_properties"]["override_enabled"],
            "values": dict(DEFAULT_CONFIG["docx_core_properties"]["values"]),
        },
    }

    if not os.path.exists(config_path):
        return config

    try:
        with open(config_path, "r", encoding="utf-8") as config_file:
            user_config = json.load(config_file)
    except Exception as exc:
        print(f"读取配置文件失败，使用默认配置: {exc}")
        return config

    metadata_keywords = user_config.get("metadata_keywords")
    if isinstance(metadata_keywords, list) and metadata_keywords:
        config["metadata_keywords"] = [str(item) for item in metadata_keywords]

    core_properties = user_config.get("docx_core_properties")
    if isinstance(core_properties, dict):
        override_enabled = core_properties.get("override_enabled")
        if isinstance(override_enabled, bool):
            config["docx_core_properties"]["override_enabled"] = override_enabled

        values = core_properties.get("values")
        if isinstance(values, dict):
            for key in config["docx_core_properties"]["values"]:
                if key in values:
                    config["docx_core_properties"]["values"][key] = str(values[key])

    return config


def iter_extents(element):
    return element.xpath('.//*[local-name()="extent"]')


def iter_drawings(container_element):
    return container_element.xpath('.//*[local-name()="drawing"]')


def iter_shapes(container_element):
    return container_element.xpath('.//*[local-name()="shape"]')


def is_tiny_element(element):
    try:
        for ext in iter_extents(element):
            width = int(ext.get("cx", 0))
            height = int(ext.get("cy", 0))
            if 0 < width < SIZE_THRESHOLD and 0 < height < SIZE_THRESHOLD:
                return True
    except Exception:
        return False
    return False


def is_target_element(element, keywords=None):
    try:
        xml_str = etree.tostring(element, encoding="unicode")
        if keywords and has_keyword_text(xml_str, keywords):
            return True
    except Exception:
        return False
    return is_tiny_element(element)


def remove_element(element, message=None):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
        if message:
            print(message)


def scrub_descr_attributes(element, keywords, location_label):
    for node in element.iter():
        for attr_name in ("descr", "title"):
            attr_value = node.get(attr_name)
            if attr_value and has_keyword_text(attr_value, keywords):
                node.set(attr_name, "")
                print(f"已定位并清除{location_label}隐藏属性 {attr_name}: {attr_value}")


def clean_header_footer_container(container, location_label):
    for drawing in iter_drawings(container._element):
        if is_target_element(drawing, HEADER_FOOTER_KEYWORDS):
            remove_element(drawing, f"已定位并清除{location_label}绘图对象水印")

    for shape in iter_shapes(container._element):
        if is_target_element(shape, HEADER_FOOTER_KEYWORDS):
            remove_element(shape, f"已定位并清除{location_label}形状对象水印")

    for paragraph in container.paragraphs:
        if has_keyword_text(paragraph.text, HEADER_FOOTER_KEYWORDS):
            print(f"已定位并清除{location_label}文本水印: {paragraph.text}")
            paragraph.text = ""


def clean_section_headers_and_footers(doc):
    for section in doc.sections:
        containers = [
            ("页眉", section.header),
            ("页脚", section.footer),
            ("首页页眉", section.first_page_header),
            ("首页页脚", section.first_page_footer),
            ("偶数页页眉", section.even_page_header),
            ("偶数页页脚", section.even_page_footer),
        ]
        for location_label, container in containers:
            clean_header_footer_container(container, location_label)


def clean_body_paragraph(paragraph, paragraph_index):
    for drawing in iter_drawings(paragraph._element):
        if is_tiny_element(drawing):
            remove_element(drawing, f"已定位并清除正文段落 {paragraph_index} 中的微小绘图水印")
        else:
            scrub_descr_attributes(drawing, BODY_DRAWING_KEYWORDS, f"正文段落 {paragraph_index} 的绘图对象")

    for shape in iter_shapes(paragraph._element):
        if is_tiny_element(shape):
            remove_element(shape, f"已定位并清除正文段落 {paragraph_index} 中的微小形状水印")
        else:
            scrub_descr_attributes(shape, BODY_DRAWING_KEYWORDS, f"正文段落 {paragraph_index} 的形状对象")

    if has_keyword_text(paragraph.text, BODY_TEXT_KEYWORDS):
        print(f"已定位并清除正文段落 {paragraph_index} 文本水印: {paragraph.text}")
        paragraph.text = ""


def clean_document_body(doc):
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        clean_body_paragraph(paragraph, index)


def clean_core_properties(doc, config):
    try:
        core_config = config["docx_core_properties"]
        keywords = config["metadata_keywords"]
        override_enabled = core_config["override_enabled"]
        configured_values = core_config["values"]
        props = doc.core_properties

        for prop_name, configured_value in configured_values.items():
            current_value = getattr(props, prop_name, "") or ""

            if override_enabled:
                if configured_value == "-":
                    continue
                if current_value != configured_value:
                    print(f"已按配置重写核心属性 {prop_name}: {current_value}")
                setattr(props, prop_name, configured_value)
                continue

            matched_keyword = first_matched_keyword(current_value, keywords)
            if matched_keyword:
                print(f"已定位并清除核心属性 {prop_name}: {current_value}，命中关键词: {matched_keyword}")
                setattr(props, prop_name, "")
    except Exception:
        pass


def clean_custom_properties_xml(xml_bytes, keywords):
    root = etree.fromstring(xml_bytes)
    for prop in list(root.findall("cp:property", CUSTOM_PROPERTY_NS)):
        prop_text = "".join(prop.itertext())
        if has_keyword_text(prop_text, keywords):
            print(f"已定位并清除自定义属性 {prop.get('name')}: {prop_text}")
            root.remove(prop)
    return etree.tostring(root, encoding="utf-8", xml_declaration=True, standalone="yes")


def rebuild_docx_without_custom_property_residue(source_docx_path, target_docx_path, keywords):
    if not os.path.exists(source_docx_path):
        return

    output_dir = os.path.dirname(os.path.abspath(target_docx_path)) or os.getcwd()
    temp_fd, temp_path = tempfile.mkstemp(suffix=".docx", dir=output_dir)
    os.close(temp_fd)

    try:
        with zipfile.ZipFile(source_docx_path, "r") as src, zipfile.ZipFile(temp_path, "w") as dst:
            for item in src.infolist():
                data = src.read(item.filename)
                if item.filename == "docProps/custom.xml":
                    data = clean_custom_properties_xml(data, keywords)
                dst.writestr(item, data)

        os.replace(temp_path, target_docx_path)
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def save_cleaned_docx(doc, output_path, metadata_keywords):
    output_dir = os.path.dirname(os.path.abspath(output_path)) or os.getcwd()
    temp_doc_fd, temp_doc_path = tempfile.mkstemp(suffix=".docx", dir=output_dir)
    os.close(temp_doc_fd)

    try:
        doc.save(temp_doc_path)
        rebuild_docx_without_custom_property_residue(temp_doc_path, output_path, metadata_keywords)
    finally:
        if os.path.exists(temp_doc_path):
            os.remove(temp_doc_path)


def clean_docx(input_path, output_path=None):
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"找不到文件: {input_path}")

    if output_path is None:
        output_path = build_output_path(input_path)

    print("正在扫描并清理 DOCX 水印...")
    config = load_config()
    doc = Document(input_path)
    clean_section_headers_and_footers(doc)
    clean_document_body(doc)
    clean_core_properties(doc, config)
    save_cleaned_docx(doc, output_path, config["metadata_keywords"])
    return output_path


if __name__ == "__main__":
    source = "水印语文试题.docx"
    target = build_output_path(source)
    result = clean_docx(source, target)
    print(f"清理成功！已生成：{result}")
